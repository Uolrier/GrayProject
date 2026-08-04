from pathlib import Path
from typing import List

from docx import Document as DocxDocument

from ..base import BaseDocumentLoader
from ..schema import Document


class WordLoader(BaseDocumentLoader):
    """
    Loader for Microsoft Word documents (.docx).
    """

    def __init__(self, file_path: str):
        self.file_path = Path(file_path)

    def load(self) -> List[Document]:
        """
        Load docx file into Document objects.
        """

        doc = DocxDocument(self.file_path)

        paragraphs = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        content = "\n".join(paragraphs)

        return [
            Document(
                page_content=content,
                metadata={
                    "source": str(self.file_path),
                    "type": "word",
                },
            )
        ]
